import fitz  # PyMuPDF
import os
import re
import smtplib
import socket
import unicodedata
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from dotenv import load_dotenv

# Load environment variables from .env file (go up 3 levels: src/python -> src -> project root)
load_dotenv(os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(__file__))), ".env"))
import json
import sqlite3
import sqlite_vss
try:
    import sqlite_vec
except ImportError:
    sqlite_vec = None
import requests
import uvicorn
import logging
import uuid
import time
import shutil
import threading
import secrets
import hashlib
import datetime
from typing import List, Dict, Optional
from fastapi import FastAPI, UploadFile, File, Form, HTTPException
from fastapi.responses import JSONResponse
from pydantic import BaseModel
from PIL import Image
import pytesseract
from docx import Document as DocxDocument
from openpyxl import load_workbook
from openai import OpenAI
try:
    from sentence_transformers import CrossEncoder
    _cross_encoder_available = True
except ImportError:
    _cross_encoder_available = False

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# Debug log file path (for PHP compatibility)
DEBUG_LOG_PATH = "/var/www/html/kapjus.kaponline.com.br/debug.log"

def debug_log(step: str, message: str, params: dict = None):
    """Write to debug.log for PHP compatibility"""
    import datetime
    timestamp = datetime.datetime.now().isoformat()
    params_str = f" | Params: {json.dumps(params)}" if params else ""
    log_entry = f"[{timestamp}] [KAPJUS-{step}] {message}{params_str}\n"
    try:
        with open(DEBUG_LOG_PATH, 'a') as f:
            f.write(log_entry)
    except Exception:
        pass  # Ignore logging errors

def log_step(step: str, message: str, params: dict = None):
    """Log to both logger and debug.log"""
    logger.info(f"[{step}] {message}")
    debug_log(step, message, params)


def send_invitation_email(to_email: str, case_id: str, magic_link: str, inviter_name: str = "Um advogado", case_name: str = "", case_description: str = "") -> bool:
    """Send invitation email via localhost SMTP (port 25, no auth)."""
    
    # Email configuration
    SMTP_HOST = os.getenv("SMTP_HOST", "localhost")
    SMTP_PORT = int(os.getenv("SMTP_PORT", "25"))
    SMTP_FROM = os.getenv("SMTP_FROM", "KapJus <noreply@janeri.com.br>")
    
    # Build case info for email
    case_info_html = ""
    case_info_text = ""
    if case_name:
        case_info_html = f"""
            <div class="case-info" style="background: #f8f9fa; border-radius: 8px; padding: 20px; margin: 20px 0;">
                <p style="margin: 0; color: #1a237e; font-size: 16px; font-weight: bold;">📁 Processo: {case_name}</p>
                {f'<p style="margin: 10px 0 0; color: #555;">{case_description}</p>' if case_description else ''}
            </div>
        """
        case_info_text = f"""
Processo: {case_name}
{f'Descrição: {case_description}' if case_description else ''}

"""
    
    # Build HTML email
    html_content = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <style>
        body {{ font-family: 'Segoe UI', Arial, sans-serif; background: #f5f5f5; margin: 0; padding: 20px; }}
        .container {{ max-width: 600px; margin: 0 auto; background: white; border-radius: 12px; overflow: hidden; box-shadow: 0 2px 10px rgba(0,0,0,0.1); }}
        .header {{ background: linear-gradient(135deg, #1a237e 0%, #283593 100%); color: white; padding: 30px; text-align: center; }}
        .header h1 {{ margin: 0; font-size: 28px; }}
        .header p {{ margin: 10px 0 0; opacity: 0.9; font-size: 14px; }}
        .content {{ padding: 30px; color: #333; line-height: 1.7; }}
        .content h2 {{ color: #1a237e; font-size: 20px; margin-top: 0; }}
        .content p {{ margin-bottom: 20px; }}
        .features {{ background: #f8f9fa; border-radius: 8px; padding: 20px; margin: 20px 0; }}
        .features ul {{ margin: 0; padding-left: 20px; }}
        .features li {{ margin: 8px 0; color: #555; }}
        .features li strong {{ color: #1a237e; }}
        .button {{ display: inline-block; background: linear-gradient(135deg, #1a237e 0%, #283593 100%); color: white; padding: 16px 32px; text-decoration: none; border-radius: 8px; font-weight: bold; font-size: 16px; margin: 20px 0; }}
        .button:hover {{ opacity: 0.9; }}
        .footer {{ background: #f5f5f5; padding: 20px; text-align: center; font-size: 12px; color: #888; }}
        .warning {{ background: #fff3cd; border-left: 4px solid #ffc107; padding: 12px; margin: 20px 0; font-size: 13px; color: #856404; }}
    </style>
</head>
<body>
    <div class="container">
        <div class="header">
            <h1>⚖️ KapJus</h1>
            <p>Sua Inteligência Artificial Jurídica</p>
        </div>
        <div class="content">
            <h2>Olá!</h2>
            <p><strong>{inviter_name}</strong> convidou você para acessar um processo jurídico no <strong>KapJus</strong>.</p>
            
            {case_info_html}
            
            <div class="features">
                <strong>O que você pode fazer no KapJus:</strong>
                <ul>
                    <li><strong>🔍 Busca Rápida</strong> - Encontre qualquer termo, pessoa ou fato em segundos</li>
                    <li><strong>📱 Acesse do Celular</strong> - Interface responsiva para usar em qualquer lugar</li>
                    <li><strong>🤖 AI Jurídica</strong> - Tire dúvidas sobre o processo com inteligência artificial</li>
                    <li><strong>📄 Análise de Documentos</strong> - Extração automática de texto e dados</li>
                </ul>
            </div>
            
            <p style="text-align: center;">
                <a href="{magic_link}" class="button">Acessar Processo</a>
            </p>
            
            <div class="warning">
                ⚠️ Este link expira em 48 horas e é único. Não compartilhe com terceiros.
            </div>
            
            <p style="font-size: 13px; color: #888;">
                Se você não solicitou este acesso, ignore este email.
            </p>
        </div>
        <div class="footer">
            KapJus © 2024 - Inteligência Artificial para o Direito<br>
            Este é um email automático, não responda.
        </div>
    </div>
</body>
</html>"""
    
    # Plain text version
    text_content = f"""KapJus - Sua Inteligência Jurídica

Olá!

{inviter_name} convidou você para acessar um processo jurídico no KapJus.

{case_info_text}O que você pode fazer:
- Busca Rápida: Encontre qualquer termo, pessoa ou fato em segundos
- Acesse do Celular: Interface responsiva para usar em qualquer lugar
- AI Jurídica: Tire dúvidas sobre o processo com inteligência artificial
- Análise de Documentos: Extração automática de texto e dados

Acesse aqui: {magic_link}

Este link expira em 48 horas e é único. Não compartilhe com terceiros.

-- 
KapJus © 2024 - Inteligência Artificial para o Direito
"""
    
    try:
        # Create message
        msg = MIMEMultipart('alternative')
        msg['Subject'] = f"Convite para acessar processo no KapJus - Caso {case_id}"
        msg['From'] = SMTP_FROM
        msg['To'] = to_email
        
        # Attach both versions
        text_part = MIMEText(text_content, 'plain', 'utf-8')
        html_part = MIMEText(html_content, 'html', 'utf-8')
        msg.attach(text_part)
        msg.attach(html_part)
        
        # Send via localhost SMTP
        with smtplib.SMTP(SMTP_HOST, SMTP_PORT) as server:
            # No authentication for localhost
            server.send_message(msg)
        
        log_step("EMAIL", f"Invitation email sent to {to_email} via {SMTP_HOST}:{SMTP_PORT}")
        return True
        
    except smtplib.SMTPException as e:
        log_step("EMAIL", f"SMTP error sending to {to_email}: {e}")
        return False
    except socket.error as e:
        log_step("EMAIL", f"Connection error to SMTP {SMTP_HOST}:{SMTP_PORT}: {e}")
        return False
    except Exception as e:
        log_step("EMAIL", f"Error sending to {to_email}: {e}")
        return False


def call_ai(prompt: str, provider: str = None) -> str:
    """Call AI with the specified or default provider. Returns the response text."""
    effective_provider = provider if provider else IA_PROVIDER
    log_step("AI", f"Calling AI with provider: {effective_provider}", {"prompt_length": len(prompt)})
    
    if effective_provider == "openrouter":
        for model in OPENROUTER_MODELS:
            try:
                response = openrouter_client.chat.completions.create(
                    model=model,
                    messages=[
                        {"role": "system", "content": "Você é um assistente jurídico helpful."},
                        {"role": "user", "content": prompt}
                    ],
                    max_tokens=2048,
                    temperature=0.7
                )
                answer = response.choices[0].message.content
                log_step("AI", f"OpenRouter response with {model}", {"length": len(answer)})
                return answer
            except Exception as e:
                log_step("AI", f"Model {model} failed, trying next", {"error": str(e)})
                continue
        raise Exception("All OpenRouter models failed")
    elif effective_provider == "gemini":
        try:
            url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-1.5-flash:generateContent?key={GEMINI_API_KEY}"
            response = requests.post(url, json={
                "contents": [{"parts": [{"text": prompt}]}]
            })
            data = response.json()
            if "candidates" in data:
                answer = data["candidates"][0]["content"]["parts"][0]["text"]
                log_step("AI", f"Gemini response", {"length": len(answer)})
                return answer
            else:
                raise Exception(f"Gemini response error: {data}")
        except Exception as e:
            log_step("AI", f"Gemini error", {"error": str(e)})
            raise e
    else:
        raise Exception(f"Unknown provider: {effective_provider}")



# Configurações de Caminhos (must be defined first)
BASE_DIR = "/var/www/html/kapjus.kaponline.com.br"
UPLOAD_DIR = os.path.join(BASE_DIR, "storage/uploads")
DB_PATH = os.path.join(BASE_DIR, "database/kapjus.db")
SOCKET_PATH = os.path.join(BASE_DIR, "socket/kapjus.sock")
CHUNKS_DIR = os.path.join(BASE_DIR, "storage/uploads/chunks")

# Chunked upload configuration
CHUNK_SIZE = 1024 * 1024  # 1MB chunks
MAX_FILE_SIZE = 100 * 1024 * 1024  # 100MB max
UPLOAD_TIMEOUT = 24 * 60 * 60  # 24 hours in seconds

# Create chunks directory on startup
os.makedirs(CHUNKS_DIR, exist_ok=True)

# Configurações de IA (Devem ser configuradas no ambiente)
OPENROUTER_API_KEY = os.getenv("OPENROUTER_API_KEY", "")
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY", "")

RETRIEVE_CHUNK_COUNT = 50
FINAL_CONTEXT_CHUNKS = 12
MIN_FINAL_CONTEXT_CHUNKS = 8

def _mask_api_key(key: str) -> str:
    """Mask sensitive key for logging."""
    if not key:
        return ""
    prefix = key[:6]
    suffix = key[-4:]
    return f"{prefix}…{suffix}"

log_step("ENV", "Environment variables loaded", {
    "openrouter_key": _mask_api_key(OPENROUTER_API_KEY),
    "gemini_key": _mask_api_key(GEMINI_API_KEY)
})
IA_PROVIDER = os.getenv("IA_PROVIDER", "gemini").lower()  # Options: gemini, openrouter

# OpenRouter Model Fallback Chain (best to worst)
OPENROUTER_MODELS_DEFAULT = [
    "deepseek/deepseek-r1:free",
    "meta-llama/llama-3.3-70b-instruct:free",
    "tngtech/deepseek-r1t2-chimera:free",
    "arcee-ai/trinity-large-preview:free",
    "qwen/qwen3-next-80b-a3b-instruct:free",
    "openai/gpt-oss-120b:free",
    "stepfun/step-3-5-flash:free",
    "z-ai/glm-4-5-air:free",
    "nvidia/nemotron-3-nano-30b-a3b:free",
    "google/gemma-3-27b:free",
    "arcee-ai/trinity-mini:free",
    "openai/gpt-oss-20b:free",
    "nvidia/nemotron-nano-9b-v2:free",
    "upstage/solar-pro-3:free",
    "venice/venice-uncensored:free",
    "liquid/lfm2.5-1.2b-instruct:free",
]

models_env = os.getenv("OPENROUTER_MODELS") or os.getenv("openrouter_models")
if models_env:
    OPENROUTER_MODELS = [model.strip() for model in models_env.split(",") if model.strip()]
else:
    OPENROUTER_MODELS = OPENROUTER_MODELS_DEFAULT

# Create OpenRouter client
openrouter_client = OpenAI(
    api_key=OPENROUTER_API_KEY,
    base_url="https://openrouter.ai/api/v1"
)

app = FastAPI()

# Enable VSS extension for vector similarity search
def enable_vss_extension(conn):
    """Load sqlite-vss or sqlite-vec extension for vector search capabilities."""
    # Try sqlite-vec first (modern replacement)
    if sqlite_vec:
        try:
            conn.enable_load_extension(True)
            sqlite_vec.load(conn)
            conn.enable_load_extension(False)
            logger.info("sqlite-vec extension loaded successfully")
            return "vec"
        except Exception as e:
            logger.warning(f"sqlite-vec not available: {e}")
    
    # Fallback to sqlite-vss
    try:
        conn.enable_load_extension(True)
        sqlite_vss.load(conn)
        conn.enable_load_extension(False)
        logger.info("sqlite-vss extension loaded successfully")
        return "vss"
    except Exception as e:
        logger.warning(f"VSS extension not available: {e}")
        return False

# Test VSS availability on module load
_vss_mode = False
try:
    _test_conn = sqlite3.connect(DB_PATH)
    _vss_mode = enable_vss_extension(_test_conn)
    _test_conn.close()
except Exception:
    pass

_vss_available = bool(_vss_mode)
logger.info(f"VSS extension available: {_vss_available} (Mode: {_vss_mode})")

# Cleanup stale uploads (older than UPLOAD_TIMEOUT)
def cleanup_stale_uploads():
    """Remove chunks directories older than UPLOAD_TIMEOUT"""
    if not os.path.exists(CHUNKS_DIR):
        return
    
    current_time = time.time()
    for upload_id in os.listdir(CHUNKS_DIR):
        upload_path = os.path.join(CHUNKS_DIR, upload_id)
        if os.path.isdir(upload_path):
            stat = os.stat(upload_path)
            if current_time - stat.st_mtime > UPLOAD_TIMEOUT:
                shutil.rmtree(upload_path)
                logger.info(f"Cleaned up stale upload: {upload_id}")

# Cleanup on startup
cleanup_stale_uploads()

class SearchQuery(BaseModel):
    case_id: str
    query: str
    top_k: int = 20
    offset: int = 0
    filename_filter: Optional[str] = None   # filtrar por nome de arquivo (substring)
    file_type_filter: Optional[str] = None  # filtrar por extensão: "pdf", "docx", "txt", etc.

def extract_text_from_pdf(pdf_path: str) -> List[Dict]:
    """
    Extrai texto de PDF. Se o PDF for escaneado (sem texto nativo),
    aplica OCR usando pytesseract.
    """
    doc = fitz.open(pdf_path)
    pages_content = []
    
    for page_num in range(len(doc)):
        page = doc.load_page(page_num)
        text = page.get_text("text")
        
        # Verificar se o texto extraído está vazio ou é muito curto (possível PDF escaneado)
        if not text or len(text.strip()) < 50:
            # Aplicar OCR na página
            pix = page.get_pixmap(matrix=fitz.Matrix(300/72, 300/72))
            img_data = pix.tobytes("png")
            try:
                from PIL import Image
                import io
                img = Image.open(io.BytesIO(img_data))
                text = pytesseract.image_to_string(img, lang='por')
            except Exception as ocr_error:
                text = f"[OCR Error: {str(ocr_error)}]"
        
        pages_content.append({"page": page_num + 1, "content": normalize_text(text)})

    return pages_content

def extract_text_from_image(image_path: str) -> List[Dict]:
    """
    Extrai texto de imagens (JPG, PNG) usando OCR com pytesseract.
    """
    try:
        img = Image.open(image_path)
        text = pytesseract.image_to_string(img, lang='por')
        return [{"page": 1, "content": normalize_text(text)}]
    except Exception as e:
        return [{"page": 1, "content": f"[OCR Error: {str(e)}]"}]

def extract_text_from_docx(docx_path: str) -> List[Dict]:
    """
    Extrai texto de documentos DOCX usando python-docx.
    """
    try:
        doc = DocxDocument(docx_path)
        full_text = []
        for paragraph in doc.paragraphs:
            full_text.append(paragraph.text)
        
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)
        
        text = "\n".join(full_text)
        return [{"page": 1, "content": normalize_text(text)}]
    except Exception as e:
        return [{"page": 1, "content": f"[DOCX Error: {str(e)}]"}]

def extract_text_from_txt(txt_path: str) -> List[Dict]:
    """
    Extrai texto de arquivos TXT.
    """
    try:
        with open(txt_path, 'r', encoding='utf-8') as f:
            text = f.read()
        return [{"page": 1, "content": normalize_text(text)}]
    except Exception as e:
        return [{"page": 1, "content": f"[TXT Error: {str(e)}]"}]

def extract_text_from_xlsx(xlsx_path: str) -> List[Dict]:
    """
    Extrai texto de arquivos Excel (XLSX) usando openpyxl.
    """
    try:
        wb = load_workbook(xlsx_path)
        full_text = []
        
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            full_text.append(f"=== Sheet: {sheet_name} ===")
            
            for row in ws.iter_rows(values_only=True):
                row_text = []
                for cell in row:
                    if cell is not None:
                        row_text.append(str(cell))
                if row_text:
                    full_text.append(" | ".join(row_text))
        
        text = "\n".join(full_text)
        return [{"page": 1, "content": text}]
    except Exception as e:
        return [{"page": 1, "content": f"[XLSX Error: {str(e)}]"}]

def extract_text_from_document(file_path: str) -> List[Dict]:
    """
    Detecta o tipo de arquivo e roteia para a função de extração apropriada.
    """
    ext = os.path.splitext(file_path)[1].lower()
    
    extractors = {
        '.pdf': extract_text_from_pdf,
        '.jpg': extract_text_from_image,
        '.jpeg': extract_text_from_image,
        '.png': extract_text_from_image,
        '.docx': extract_text_from_docx,
        '.txt': extract_text_from_txt,
        '.xlsx': extract_text_from_xlsx,
    }
    
    if ext in extractors:
        return extractors[ext](file_path)
    else:
        return [{"page": 1, "content": f"[Unsupported format: {ext}]"}]


def is_section_heading(line: str) -> bool:
    """
    Determina se uma linha em maiúsculas representa um título/seção lógica.
    """
    stripped = line.strip()
    if len(stripped) < 5:
        return False

    letters = sum(ch.isalpha() for ch in stripped)
    if letters < 4:
        return False

    uppercase_letters = sum(1 for ch in stripped if ch.isalpha() and ch.isupper())
    if uppercase_letters < 3:
        return False

    ratio = uppercase_letters / letters if letters else 0
    return ratio >= 0.7


def split_page_into_sections(page_content: str, page_number: int) -> List[Dict]:
    """
    Divide o texto de uma página em blocos baseados em títulos/sessões em maiúsculas.
    """
    sections = []
    buffer = []
    current_heading = None

    for line in page_content.splitlines():
        stripped_line = line.strip()
        if is_section_heading(stripped_line):
            if buffer:
                chunk_text = "\n".join(buffer).strip()
                if chunk_text:
                    sections.append({"page": page_number, "content": chunk_text, "heading": current_heading})
            current_heading = stripped_line
            buffer = [stripped_line]
            continue

        buffer.append(line)

    if buffer:
        chunk_text = "\n".join(buffer).strip()
        if chunk_text:
            sections.append({"page": page_number, "content": chunk_text, "heading": current_heading})

    return sections


def prepare_logical_chunks(pages: List[Dict]) -> List[Dict]:
    """
    Agrupa todas as páginas em chunks lógicos (títulos em maiúsculas, se disponíveis).
    """
    logical_chunks = []
    for page in pages:
        content = page.get("content", "")
        if not content.strip():
            continue

        sections = split_page_into_sections(content, page["page"])
        logical_chunks.extend(sections or [{"page": page["page"], "content": content.strip(), "heading": None}])

    return logical_chunks


def normalize_text(text: str) -> str:
    """Normaliza texto para Unicode NFC, garantindo consistência de diacríticos."""
    if not text:
        return text
    return unicodedata.normalize('NFC', text)


def sanitize_for_fts(term: str) -> Optional[str]:
    """
    Remove caracteres inválidos e normaliza termos para consultas FTS5.
    """
    if not term or not term.strip():
        return None

    cleaned = re.sub(r"[^\wÀ-ÖØ-öø-ÿ]+", " ", term, flags=re.UNICODE).strip()
    cleaned = re.sub(r"\s+", " ", cleaned)

    return cleaned.lower() if cleaned else None


def store_document(case_id: str, filename: str, pages: List[Dict]):
    """
    Armezena o conteúdo extraído no banco de dados SQLite.
    """
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    cursor.execute("CREATE TABLE IF NOT EXISTS documents (id INTEGER PRIMARY KEY AUTOINCREMENT, case_id TEXT, filename TEXT, page_number INTEGER, content TEXT)")
    cursor.execute("CREATE VIRTUAL TABLE IF NOT EXISTS documents_fts USING fts5(content, case_id UNINDEXED, filename UNINDEXED, page_number UNINDEXED)")

    chunks = prepare_logical_chunks(pages)
    logger.info(f"Storing {filename} ({len(chunks)} logical chunks)")

    for chunk in chunks:
        chunk_content = chunk["content"]
        cursor.execute("INSERT INTO documents (case_id, filename, page_number, content) VALUES (?, ?, ?, ?)", (case_id, filename, chunk['page'], chunk_content))
        doc_id = cursor.lastrowid
        cursor.execute("INSERT INTO documents_fts (rowid, content, case_id, filename, page_number) VALUES (?, ?, ?, ?, ?)", (doc_id, chunk_content, case_id, filename, chunk['page']))
    
    conn.commit()
    conn.close()

# ==================== VSS (Vector Similarity Search) Infrastructure ====================

def create_vss_table(conn):
    """Create VSS virtual table for vector similarity search."""
    if not _vss_available:
        logger.warning("VSS extension not available, skipping table creation")
        return False
    
    cursor = conn.cursor()
    try:
        if _vss_mode == "vec":
            cursor.execute("""
                CREATE VIRTUAL TABLE IF NOT EXISTS vec_chunks USING vec0(
                    rowid INTEGER PRIMARY KEY,
                    embedding FLOAT[1536]
                )
            """)
            logger.info("sqlite-vec vec_chunks virtual table created/verified")
        else:
            cursor.execute("""
                CREATE VIRTUAL TABLE IF NOT EXISTS vss_chunks USING vss0(
                    embedding(1536)  -- OpenAI text-embedding-3-small dimension
                )
            """)
            logger.info("sqlite-vss vss_chunks virtual table created/verified")
        conn.commit()
        return True
    except Exception as e:
        logger.error(f"Failed to create vector table: {e}")
        return False

def get_embedding(text: str, model: str = "text-embedding-3-small") -> Optional[List[float]]:
    """Generate embedding vector for given text using OpenAI API via OpenRouter."""
    if not text or not text.strip():
        return None
    
    try:
        response = openrouter_client.embeddings.create(
            model=model,
            input=text[:8000]  # Truncate to avoid token limits
        )
        embedding = response.data[0].embedding
        logger.debug(f"Generated embedding for text ({len(text)} chars) using {model}")
        return embedding
    except Exception as e:
        logger.error(f"Failed to generate embedding: {e}")
        return None

# Table name based on available extension
VEC_TABLE = "vec_chunks" if _vss_mode == "vec" else "vss_chunks"

def store_document_with_embedding(case_id: str, filename: str, pages: List[Dict]):
    """Store document with both text and vector embeddings."""
    conn = sqlite3.connect(DB_PATH)
    
    if _vss_available:
        create_vss_table(conn)
        enable_vss_extension(conn) # Ensure extension is loaded for this connection
    
    cursor = conn.cursor()
    cursor.execute("CREATE TABLE IF NOT EXISTS documents (id INTEGER PRIMARY KEY AUTOINCREMENT, case_id TEXT, filename TEXT, page_number INTEGER, content TEXT)")
    cursor.execute("CREATE VIRTUAL TABLE IF NOT EXISTS documents_fts USING fts5(content, case_id UNINDEXED, filename UNINDEXED, page_number UNINDEXED)")

    chunks = prepare_logical_chunks(pages)
    logger.info(f"Storing embeddings for {filename} ({len(chunks)} logical chunks)")

    for chunk in chunks:
        content = chunk['content']
        page_num = chunk['page']
        
        # Store text in documents and FTS tables
        cursor.execute("INSERT INTO documents (case_id, filename, page_number, content) VALUES (?, ?, ?, ?)", 
                      (case_id, filename, page_num, content))
        doc_id = cursor.lastrowid
        cursor.execute("INSERT INTO documents_fts (rowid, content, case_id, filename, page_number) VALUES (?, ?, ?, ?, ?)", 
                      (doc_id, content, case_id, filename, page_num))
        
        # Generate and store embedding if VSS is available
        if _vss_available:
            embedding = get_embedding(content)
            if embedding:
                try:
                    if _vss_mode == "vec":
                        # sqlite-vec needs the embedding as a blob/bytes or JSON string depending on how it was created
                        # For float[1536], it accepts bytes
                        import array
                        embedding_bytes = array.array('f', embedding).tobytes()
                        cursor.execute(f"INSERT INTO {VEC_TABLE} (rowid, embedding) VALUES (?, ?)", 
                                      (doc_id, embedding_bytes))
                    else:
                        cursor.execute(f"INSERT INTO {VEC_TABLE} (rowid, embedding) VALUES (?, ?)", 
                                      (doc_id, json.dumps(embedding)))
                    logger.debug(f"Stored embedding for {filename} page {page_num}")
                except Exception as e:
                    logger.warning(f"Failed to store embedding for {filename} page {page_num}: {e}")
    
    conn.commit()
    conn.close()
    logger.info(f"Stored document with embeddings: {filename} ({len(pages)} pages)")

def index_document_embeddings(conn: sqlite3.Connection, case_id: str = None):
    """Generate and store embeddings for documents without vectors."""
    if not _vss_available:
        logger.warning("VSS extension not available, skipping embedding indexing")
        return 0
    
    # Ensure extension is loaded for this connection
    enable_vss_extension(conn)
    
    cursor = conn.cursor()
    
    # Create vss_chunks table if it doesn't exist
    create_vss_table(conn)
    
    # Find documents without embeddings
    if case_id:
        cursor.execute(f"""
            SELECT d.id, d.case_id, d.filename, d.page_number, d.content
            FROM documents d
            LEFT JOIN {VEC_TABLE} v ON d.id = v.rowid
            WHERE v.rowid IS NULL AND d.case_id = ?
            ORDER BY d.id
        """, (case_id,))
    else:
        cursor.execute(f"""
            SELECT d.id, d.case_id, d.filename, d.page_number, d.content
            FROM documents d
            LEFT JOIN {VEC_TABLE} v ON d.id = v.rowid
            WHERE v.rowid IS NULL
            ORDER BY d.id
        """)
    
    docs_to_index = cursor.fetchall()
    
    if not docs_to_index:
        logger.info("No documents need embedding indexing")
        return 0
    
    logger.info(f"Indexing {len(docs_to_index)} documents for vector search using {VEC_TABLE}")
    indexed_count = 0
    
    for doc in docs_to_index:
        doc_id, doc_case_id, filename, page_num, content = doc
        embedding = get_embedding(content)
        if embedding:
            try:
                if _vss_mode == "vec":
                    import array
                    embedding_bytes = array.array('f', embedding).tobytes()
                    cursor.execute(f"INSERT INTO {VEC_TABLE} (rowid, embedding) VALUES (?, ?)", 
                                  (doc_id, embedding_bytes))
                else:
                    cursor.execute(f"INSERT INTO {VEC_TABLE} (rowid, embedding) VALUES (?, ?)", 
                                  (doc_id, json.dumps(embedding)))
                indexed_count += 1
                logger.debug(f"Indexed embedding for {filename} page {page_num}")
            except Exception as e:
                logger.warning(f"Failed to index embedding for {filename} page {page_num}: {e}")
    
    conn.commit()
    logger.info(f"Indexed {indexed_count} embeddings successfully")
    return indexed_count

class SetupVSSQuery(BaseModel):
    case_id: Optional[str] = None

@app.post("/setup_vss")
async def setup_vss(query: SetupVSSQuery):
    """Setup VSS infrastructure: create vector table and index existing documents."""
    if not _vss_available:
        return JSONResponse(
            status_code=503,
            content={"status": "error", "message": "VSS/Vector extension not available"}
        )
    
    conn = sqlite3.connect(DB_PATH)
    
    try:
        # Load extension for this connection
        enable_vss_extension(conn)

        # Create table
        table_created = create_vss_table(conn)
        
        if not table_created:
            raise Exception(f"Failed to create {VEC_TABLE} table")
        
        # Index existing documents
        indexed = index_document_embeddings(conn, query.case_id)
        
        # Get stats
        cursor = conn.cursor()
        cursor.execute(f"SELECT COUNT(*) FROM {VEC_TABLE}")
        total_embeddings = cursor.fetchone()[0]
        
        return {
            "status": "success",
            "message": "Vector infrastructure setup complete",
            "vss_available": _vss_available,
            "vss_mode": _vss_mode,
            "table_name": VEC_TABLE,
            "indexed_documents": indexed,
            "total_embeddings": total_embeddings
        }
    
    except Exception as e:
        logger.error(f"Vector setup failed: {e}")
        return JSONResponse(
            status_code=500,
            content={"status": "error", "message": str(e)}
        )
    finally:
        conn.close()

@app.get("/vss_status")
async def vss_status():
    """Get VSS/Vector infrastructure status."""
    conn = sqlite3.connect(DB_PATH)
    enable_vss_extension(conn)
    cursor = conn.cursor()
    
    status = {
        "vss_available": _vss_available,
        "vss_mode": _vss_mode,
        "table_name": VEC_TABLE,
        "vector_table_exists": False,
        "total_embeddings": 0,
        "documents_without_embeddings": 0
    }
    
    try:
        if _vss_available:
            cursor.execute(f"SELECT name FROM sqlite_master WHERE type='table' AND name='{VEC_TABLE}'")
            if cursor.fetchone():
                status["vector_table_exists"] = True
                cursor.execute(f"SELECT COUNT(*) FROM {VEC_TABLE}")
                status["total_embeddings"] = cursor.fetchone()[0]
                
                cursor.execute(f"""
                    SELECT COUNT(*) FROM documents d
                    LEFT JOIN {VEC_TABLE} v ON d.id = v.rowid
                    WHERE v.rowid IS NULL
                """)
                status["documents_without_embeddings"] = cursor.fetchone()[0]
    except Exception as e:
        logger.error(f"Failed to get vector status: {e}")
    finally:
        conn.close()
    
    return status

# ==================== HYBRID SEARCH PHASE 2 ====================

_INTERROGATIVE_PATTERN = re.compile(
    r'\b(como|quando|qual|quais|quem|onde|por que|porque|o que|quanto|quantos|quantas)\b',
    re.IGNORECASE | re.UNICODE,
)

_COMMAND_PATTERN = re.compile(
    r'\b(resuma|liste|explique|descreva|quais|detalhe|aponte|fale sobre)\b',
    re.IGNORECASE | re.UNICODE,
)

def _is_simple_query(question: str) -> bool:
    """
    Retorna True se a query é simples (≤ 3 tokens, sem palavras interrogativas, sem comandos e sem '?').
    Queries simples não se beneficiam de expansão via LLM — apenas adicionam latência.
    """
    tokens = question.strip().split()
    if len(tokens) > 3:
        return False
    if '?' in question:
        return False
    if _INTERROGATIVE_PATTERN.search(question):
        return False
    if _COMMAND_PATTERN.search(question):
        return False
    return True


def expand_query(question: str, provider: str = None) -> Dict[str, any]:
    """
    Expand user question into keywords for FTS5 and semantic query for VSS.
    Uses AI to generate optimized search terms.
    For short, simple queries (≤ 2 tokens, no interrogative words) skips the LLM call.
    """
    log_step("EXPAND", f"Iniciando expansao da pergunta", {"question": question[:200]})

    if not question or not question.strip():
        log_step("EXPAND", "Pergunta vazia recebida", {"keywords": [], "semantic_query": ""})
        return {"keywords": [], "semantic_query": ""}

    if _is_simple_query(question):
        log_step("EXPAND", "Query simples — expansao via LLM ignorada", {"question": question})
        return {
            "keywords": question.strip().split(),
            "semantic_query": question.strip(),
        }
    
    prompt = f"""###
Você é um especialista em Recuperação de Informação (IR) aplicado ao Direito e Perícia Digital. Analise a pergunta do usuário e retorne um JSON estruturado para uma busca híbrida (FTS5 + Vetorial).

Instruções para Keywords (FTS5):
- Inclua termos técnicos (jurídicos e forenses) relacionados ao tema.
- Adicione variações de gênero e número, além de sinônimos verbais e substantivos.
- Inclua termos que costumam aparecer em cabeçalhos ou rodapés de documentos similares (ex: 'conclui-se', 'ante o exposto').

Instruções para Semantic Query:
- Transforme a dúvida em uma afirmação técnica densa que descreva o conteúdo esperado na resposta, simulando a escrita de um perito ou magistrado.

Pergunta do Usuário: '{question}'

Responda APENAS o JSON, sem markdown ou textos explicativos: {{ "keywords": [], "semantic_query": "" }}
"""
    
    log_step("EXPAND", f"Prompt gerado para IA", {"prompt": prompt[:300]})
    
    try:
        content = call_ai(prompt, provider)
        log_step("EXPAND", f"Resposta raw da IA", {"content": content[:500]})
        
        # Clean up if model adds markdown formatting
        if content.startswith("```"):
            content = "\n".join(content.split("\n")[1:-1])
        
        result = json.loads(content)
        keywords = result.get('keywords', [])
        semantic_query = result.get('semantic_query', '')
        
        log_step("EXPAND", f"Expansao concluida", {"keywords": keywords, "semantic_query": semantic_query[:200]})
        
        return {
            "keywords": keywords,
            "semantic_query": semantic_query if semantic_query else question
        }
    except Exception as e:
        log_step("EXPAND", f"Erro na expansao, usando fallback", {"error": str(e)})
        # Fallback: use question as-is
        fallback_keywords = question.split()[:10]
        return {
            "keywords": fallback_keywords,
            "semantic_query": question
        }

def hybrid_search(keywords: List[str], semantic_query: str, case_id: Optional[str] = None, limit: int = 5, original_question: str = None) -> Dict[str, List[Dict]]:
    """
    Execute parallel hybrid search using FTS5 and VSS.
    Returns dict with 'fts_results' and 'vss_results'.
    """
    log_step("HYBRID", f"Iniciando busca hibrida", {"keywords": keywords, "semantic_query": semantic_query[:200], "case_id": case_id, "limit": limit})
    
    fts_results = []
    vss_results = []
    
    # Sanitize keywords before building FTS queries
    sanitized_keywords = []
    for term in keywords:
        sanitized = sanitize_for_fts(term)
        if sanitized:
            sanitized_keywords.append(sanitized)

    original_terms = []
    if original_question:
        for term in original_question.split():
            sanitized = sanitize_for_fts(term)
            if sanitized and len(sanitized) > 2:
                original_terms.append(sanitized)

    all_keywords = original_terms + sanitized_keywords
    seen = set()
    all_keywords = [x for x in all_keywords if not (x in seen or seen.add(x))]

    log_step("HYBRID", "Sanitized keywords for FTS5", {"sanitized": sanitized_keywords, "original_terms": original_terms})
    
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    
    # FTS5 Query
    try:
        if all_keywords:
            # Build FTS5 query with OR operator for broader matching
            # Each term with prefix matching for flexibility
            fts_query_terms = [term + "*" for term in all_keywords]
            fts_query = " OR ".join(fts_query_terms)
            log_step("HYBRID-FTS5", f"Executando busca FTS5", {"query": fts_query, "original_terms": original_terms if original_question else []})
            
            if case_id:
                cursor.execute("""
                    SELECT rowid, content, filename, page_number,
                           bm25(documents_fts) as rank
                    FROM documents_fts 
                    WHERE case_id = ? AND documents_fts MATCH ?
                    ORDER BY rank
                    LIMIT ?
                """, (case_id, fts_query, limit))
            else:
                cursor.execute("""
                    SELECT rowid, content, filename, page_number,
                           bm25(documents_fts) as rank
                    FROM documents_fts 
                    WHERE documents_fts MATCH ?
                    ORDER BY rank
                    LIMIT ?
                """, (fts_query, limit))
            
            for row in cursor.fetchall():
                fts_results.append({
                    "rowid": row[0],
                    "content": row[1],
                    "filename": row[2],
                    "page": row[3],
                    "rank": row[4]
                })
        
        log_step("HYBRID-FTS5", f"Resultados FTS5 encontrados", {"count": len(fts_results)})
    except Exception as e:
        log_step("HYBRID-FTS5", f"Erro na busca FTS5", {"error": str(e)})
    
    # VSS Query (Vector Similarity Search)
    try:
        if _vss_available and semantic_query:
            log_step("HYBRID-VSS", f"Gerando embedding para busca vetorial", {"query": semantic_query[:200]})
            embedding = get_embedding(semantic_query)
            
            if embedding:
                log_step("HYBRID-VSS", f"Embedding gerado", {"dimension": len(embedding)})
                enable_vss_extension(conn) # Ensure loaded
                
                if _vss_mode == "vec":
                    import array
                    embedding_blob = array.array('f', embedding).tobytes()
                    
                    if case_id:
                        log_step("HYBRID-VSS", f"Executando sqlite-vec com filtro case_id", {"case_id": case_id})
                        cursor.execute(f"""
                            SELECT v.rowid, d.content, d.filename, d.page_number,
                                   vec_distance_cosine(v.embedding, ?) as distance
                            FROM {VEC_TABLE} v
                            JOIN documents d ON v.rowid = d.id
                            WHERE d.case_id = ?
                            ORDER BY distance
                            LIMIT ?
                        """, (embedding_blob, case_id, limit))
                    else:
                        log_step("HYBRID-VSS", "Executando sqlite-vec sem filtro")
                        cursor.execute(f"""
                            SELECT v.rowid, d.content, d.filename, d.page_number,
                                   vec_distance_cosine(v.embedding, ?) as distance
                            FROM {VEC_TABLE} v
                            JOIN documents d ON v.rowid = d.id
                            ORDER BY distance
                            LIMIT ?
                        """, (embedding_blob, limit))
                else:
                    # Original sqlite-vss logic
                    if case_id:
                        log_step("HYBRID-VSS", f"Executando VSS com filtro case_id", {"case_id": case_id})
                        cursor.execute("""
                            SELECT v.rowid, d.content, d.filename, d.page_number,
                                   vss_distance(v.embedding, ?)
                            FROM vss_chunks v
                            JOIN documents d ON v.rowid = d.id
                            WHERE d.case_id = ?
                            ORDER BY vss_distance(v.embedding, ?)
                            LIMIT ?
                        """, (json.dumps(embedding), case_id, json.dumps(embedding), limit))
                    else:
                        log_step("HYBRID-VSS", "Executando VSS sem filtro")
                        cursor.execute("""
                            SELECT v.rowid, d.content, d.filename, d.page_number,
                                   vss_distance(v.embedding, ?)
                            FROM vss_chunks v
                            JOIN documents d ON v.rowid = d.id
                            ORDER BY vss_distance(v.embedding, ?)
                            LIMIT ?
                        """, (json.dumps(embedding), json.dumps(embedding), limit))
                
                for row in cursor.fetchall():
                    vss_results.append({
                        "rowid": row[0],
                        "content": row[1],
                        "filename": row[2],
                        "page": row[3],
                        "distance": row[4]
                    })
        else:
            log_step("HYBRID-VSS", f"VSS pulado", {"vss_available": _vss_available, "has_query": bool(semantic_query)})
        
        log_step("HYBRID-VSS", f"Resultados VSS encontrados", {"count": len(vss_results)})
    except Exception as e:
        log_step("HYBRID-VSS", f"Erro na busca VSS", {"error": str(e)})
    
    conn.close()
    
    log_step("HYBRID", f"Busca hibrida concluida", {"fts_count": len(fts_results), "vss_count": len(vss_results)})
    
    return {
        "fts_results": fts_results,
        "vss_results": vss_results
    }

def rank_results(fts_results: List[Dict], vss_results: List[Dict], top_k: int = 5, case_id: str = None) -> List[Dict]:
    """
    Merge and rank results using Reciprocal Rank Fusion (RRF).
    score = 1 / (rank * 60 + position)
    If case_id is provided, validate that all rowids belong to it.
    """
    log_step("RANK", f"Iniciando ranking de resultados", {"fts_count": len(fts_results), "vss_count": len(vss_results), "top_k": top_k, "case_id": case_id})
    
    if not fts_results and not vss_results:
        log_step("RANK", "Sem resultados para ranquear")
        return []
    
    # Validate rowids against case_id if provided
    if case_id:
        conn = sqlite3.connect(DB_PATH)
        cursor = conn.cursor()
        cursor.execute("SELECT id FROM documents WHERE case_id = ?", (str(case_id),))
        valid_rowids = set(row[0] for row in cursor.fetchall())
        conn.close()
        
        fts_results = [r for r in fts_results if r.get("rowid") in valid_rowids]
        vss_results = [r for r in vss_results if r.get("rowid") in valid_rowids]
        log_step("RANK", f"Apos validacao case_id: fts={len(fts_results)}, vss={len(vss_results)}")
    
    ranked_items = {}
    
    # Process FTS5 results with RRF
    for position, item in enumerate(fts_results):
        rowid = item["rowid"]
        rank = position + 1
        rrf_score = 1.0 / (60 + rank)
        
        if rowid not in ranked_items:
            ranked_items[rowid] = {"rowid": rowid, "content": item["content"], 
                                   "filename": item["filename"], "page": item["page"],
                                   "case_id": case_id,
                                   "fts_rank": rank, "vss_rank": None, "rrf_score": rrf_score}
        else:
            ranked_items[rowid]["rrf_score"] += rrf_score
            ranked_items[rowid]["fts_rank"] = min(ranked_items[rowid].get("fts_rank", 999), rank)
    
    log_step("RANK", f"Apos FTS5: {len(ranked_items)} items unicos")
    
    # Process VSS results with RRF
    for position, item in enumerate(vss_results):
        rowid = item["rowid"]
        rank = position + 1
        rrf_score = 1.0 / (60 + rank)
        
        if rowid not in ranked_items:
            ranked_items[rowid] = {"rowid": rowid, "content": item["content"], 
                                   "filename": item["filename"], "page": item["page"],
                                   "case_id": case_id,
                                   "fts_rank": None, "vss_rank": rank, "rrf_score": rrf_score}
        else:
            ranked_items[rowid]["rrf_score"] += rrf_score
            ranked_items[rowid]["vss_rank"] = min(ranked_items[rowid].get("vss_rank", 999), rank)
    
    log_step("RANK", f"Apos VSS: {len(ranked_items)} items unicos")
    
    # Sort by combined RRF score and return top-k
    sorted_results = sorted(ranked_items.values(), key=lambda x: x["rrf_score"], reverse=True)
    
    log_step("RANK", f"Resultados ranqueados", {"total": len(sorted_results)})
    for i, item in enumerate(sorted_results[:10]):
        log_step("RANK", f"  Pos {i+1}: rowid={item['rowid']}, score={item['rrf_score']:.6f}, file={item['filename']}:{item['page']}")
    
    final_results = sorted_results[:top_k]
    log_step("RANK", f"Top-K final", {"count": len(final_results)})
    
    return final_results


def select_final_context_chunks(question: str, candidates: List[Dict], case_id: str = None) -> List[Dict]:
    """Re-rank candidates via cross-encoder and trim/pad to final context size."""
    log_step("DEBUG_SELECT", f"ENTRADA candidates={len(candidates)}, case_id={case_id}")
    
    # Filter candidates by case_id if provided
    if case_id:
        original_count = len(candidates)
        candidates = [c for c in candidates if c.get("case_id") == case_id]
        log_step("DEBUG_SELECT", f"APOS FILTRO case_id: {original_count} -> {len(candidates)}")
        if len(candidates) == 0:
            log_step("DEBUG_SELECT", "CANDIDATES VAZIOS APÓS FILTRO - TODOS PERDERAM case_id!")
    
    log_step("DEBUG_SELECT", f"ANTES cross_encoder: {len(candidates)} candidates")
    final_chunks = cross_encoder_rerank(question, candidates, FINAL_CONTEXT_CHUNKS)
    log_step("DEBUG_SELECT", f"APOS cross_encoder: {len(final_chunks)} chunks")

    if len(final_chunks) < MIN_FINAL_CONTEXT_CHUNKS:
        needed = MIN_FINAL_CONTEXT_CHUNKS - len(final_chunks)
        for chunk in candidates:
            if chunk not in final_chunks:
                # Only add chunks from the correct case_id
                if case_id is None or chunk.get("case_id") == case_id:
                    final_chunks.append(chunk)
                    needed -= 1
                    if needed == 0:
                        break

    return final_chunks[:FINAL_CONTEXT_CHUNKS]


_local_cross_encoder = None

def get_cross_encoder():
    """Lazy-load the local cross-encoder model."""
    global _local_cross_encoder
    if not _cross_encoder_available:
        return None
    if _local_cross_encoder is None:
        try:
            log_step("CROSS", "Loading local cross-encoder model...")
            # ms-marco-MiniLM-L-6-v2 is a good balance of speed and quality
            _local_cross_encoder = CrossEncoder('cross-encoder/ms-marco-MiniLM-L-6-v2', max_length=512)
            log_step("CROSS", "Local cross-encoder loaded successfully")
        except Exception as e:
            log_step("CROSS", f"Failed to load local cross-encoder: {e}")
            return None
    return _local_cross_encoder

def cross_encoder_rerank(question: str, candidates: List[Dict], top_k: int) -> List[Dict]:
    """
    Use a cross-encoder to select the most relevant chunks.
    Tries local model first, then falls back to AI (OpenRouter/Gemini).
    """
    if not candidates:
        return []

    # Try local cross-encoder first
    local_ce = get_cross_encoder()
    if local_ce:
        try:
            log_step("RERANK", f"Using local cross-encoder for {len(candidates)} candidates")
            start_time = time.time()
            
            # Prepare pairs (question, chunk_content)
            pairs = [[question, c.get("content", "")[:1000]] for c in candidates]
            scores = local_ce.predict(pairs)
            
            # Zip and sort by score
            scored_candidates = sorted(zip(candidates, scores), key=lambda x: x[1], reverse=True)
            
            # Log results
            duration = time.time() - start_time
            log_step("RERANK", f"Local rerank finished in {duration:.2f}s")
            
            return [c for c, score in scored_candidates[:top_k]]
        except Exception as e:
            log_step("RERANK", f"Local cross-encoder failed: {e}")
            # Fall back to AI-based reranking

    # Original AI-based reranking logic
    prompt_chunks = []
    for idx, chunk in enumerate(candidates):
        snippet = chunk.get("content", "").replace("\n", " ")[:400]
        prompt_chunks.append({
            "idx": idx,
            "source": f"{chunk.get('filename', 'unknown')} (pg {chunk.get('page', '?')})",
            "excerpt": snippet
        })

    prompt = f"""Você é um cross-encoder jurídico. Dada a pergunta '{question}', ordene os trechos abaixo por relevância.
Retorne apenas JSON no formato {{ "order": [índices], "scores": {{ "índice": nota, ... }} }} onde 'order' é a ordem decrescente de relevância
e cada índice corresponde ao chunk original explicado em `prompt_chunks`.

Chunks:
{json.dumps(prompt_chunks, ensure_ascii=False)}
"""

    try:
        response = call_ai(prompt)
        content = response.strip()
        if content.startswith("```"):
            content = "\n".join(content.split("\n")[1:-1])
        data = json.loads(content)
        order = data.get("order", [])
    except Exception as e:
        log_step("RERANK", "Cross-encoder falhou, usando ordem original", {"error": str(e)})
        # Fallback: retornar candidatos na ordem original se order estiver vazio
        return candidates[:top_k]

    ordered = []
    for idx in order:
        if isinstance(idx, int) and 0 <= idx < len(candidates):
            ordered.append(candidates[idx])
        if len(ordered) == top_k:
            break

    if len(ordered) < top_k:
        for chunk in candidates:
            if chunk not in ordered:
                ordered.append(chunk)
            if len(ordered) == top_k:
                break

    log_step("RERANK", "Cross-encoder selected chunks", {"selected": len(ordered), "top_k": top_k})
    return ordered

def generate_response_with_context(question: str, ranked_chunks: List[Dict], provider: str = None) -> str:
    """
    Generate AI response using the ranked context chunks.
    """
    log_step("RESPONSE", f"Iniciando geracao de resposta", {"question": question[:200], "chunks_count": len(ranked_chunks)})
    
    if not ranked_chunks:
        log_step("RESPONSE", "Sem chunks para gerar resposta")
        return "Não foi possível encontrar informações relevantes para responder à pergunta."
    
    # Build context from chunks
    context_parts = []
    for i, chunk in enumerate(ranked_chunks, 1):
        source_info = f"[{chunk.get('filename', 'unknown')}:pg {chunk.get('page', '?')}]"
        content = chunk.get('content', '')[:1000]  # Limit chunk size
        context_parts.append(f"--- Fonte {i} {source_info} ---\n{content}")
    
    context = "\n\n".join(context_parts)
    log_step("RESPONSE", f"Contexto construido", {"chunks": len(context_parts), "context_length": len(context)})
    
    prompt = f"""Você é um assistente técnico jurídico. Abaixo estão trechos recuperados de documentos. 
Utilize-os estritamente para responder à pergunta final. Se houver contradição, priorize os trechos mais relevantes.

CONTEXTO RECUPERADO:
{context}

PERGUNTA: > {question}

Resposta (cite as fontes utilizadas):"""
    
    try:
        answer = call_ai(prompt, provider)
        log_step("RESPONSE", f"Resposta gerada pela IA", {"answer_length": len(answer), "answer_preview": answer[:200]})
        
        return answer
    except Exception as e:
        log_step("RESPONSE", f"Erro ao gerar resposta", {"error": str(e)})
        return f"Erro ao gerar resposta: {str(e)}"

class HybridSearchQuery(BaseModel):
    question: str
    case_id: Optional[str] = None
    top_k: int = 5

@app.post("/hybrid_search")
async def hybrid_search_endpoint(query: HybridSearchQuery):
    """
    Hybrid search endpoint combining FTS5 and VSS search with AI response generation.
    
    Flow:
    1. Expand question using AI
    2. Parallel FTS5 + VSS search
    3. RRF ranking of results
    4. Generate response with context
    """
    logger.info(f"Hybrid search request: {query.question[:100]}...")
    
    try:
        # Step 1: Expand query
        expanded = expand_query(query.question)
        
        # Step 2: Parallel hybrid search
        search_results = hybrid_search(
            keywords=expanded["keywords"],
            semantic_query=expanded["semantic_query"],
            case_id=query.case_id,
            limit=RETRIEVE_CHUNK_COUNT,
            original_question=query.question
        )

        # Step 3: Rank results
        ranked_candidates = rank_results(
            fts_results=search_results["fts_results"],
            vss_results=search_results["vss_results"],
            top_k=RETRIEVE_CHUNK_COUNT,
            case_id=query.case_id
        )

        final_chunks = select_final_context_chunks(query.question, ranked_candidates, query.case_id)

        # Step 4: Generate response
        answer = generate_response_with_context(query.question, final_chunks)

        # Build sources list
        sources = [
            {
                "rowid": chunk["rowid"],
                "filename": chunk["filename"],
                "page": chunk["page"],
                "score": round(chunk["rrf_score"], 6)
            }
            for chunk in final_chunks
        ]
        
        return {
            "answer": answer,
            "sources": sources,
            "query_expansion": {
                "keywords": expanded["keywords"],
                "semantic_query": expanded["semantic_query"]
            },
            "stats": {
                "fts_results": len(search_results["fts_results"]),
                "vss_results": len(search_results["vss_results"]),
                "final_results": len(final_chunks)
            }
        }
    except Exception as e:
        logger.error(f"Hybrid search failed: {e}")
        return JSONResponse(
            status_code=500,
            content={"error": str(e)}
        )

@app.post("/ask_ia")
async def ask_ia_v2(
    case_id: str = Form(...), 
    question: str = Form(...), 
    provider: str = Form(None),
    use_hybrid: bool = Form(True)  # DEFAULT TO HYBRID
):
    """
    Unified endpoint for AI questions using hybrid search pipeline by default.
    
    Flow:
    1. Query expansion (keywords + semantic_query) via AI
    2. Parallel FTS5 + VSS search
    3. RRF ranking
    4. Response generation with context
    
    Args:
        case_id: ID do caso/processual
        question: Pergunta do usuário
        provider: Provedor de IA (openrouter ou gemini). Se não informado, usa IA_PROVIDER do .env
        use_hybrid: Se True (default), usa hybrid search; se False, usa FTS5 puro
    """
    # Use provider from request, or fall back to environment variable
    effective_provider = provider if provider else IA_PROVIDER
    log_step("ASK_IA", f"Requisicao recebida", {"case_id": case_id, "question": question[:100], "provider": effective_provider, "use_hybrid": use_hybrid})
    
    if use_hybrid:
        log_step("ASK_IA", "Modo: HYBRID search pipeline")
        
        # Step 1: Expand query
        expanded = expand_query(question, effective_provider)
        log_step("ASK_IA", f"Expansao: keywords={expanded['keywords']}, semantic={expanded['semantic_query'][:50]}...")
        
        # Step 2: Parallel hybrid search
        search_results = hybrid_search(
            keywords=expanded["keywords"],
            semantic_query=expanded["semantic_query"],
            case_id=case_id,
            limit=RETRIEVE_CHUNK_COUNT,
            original_question=question
        )
        log_step("ASK_IA", f"Busca: FTS5={len(search_results['fts_results'])}, VSS={len(search_results['vss_results'])}")
        
        # Step 3: Rank results with RRF
        ranked_candidates = rank_results(
            fts_results=search_results["fts_results"],
            vss_results=search_results["vss_results"],
            top_k=RETRIEVE_CHUNK_COUNT,
            case_id=case_id
        )
        final_chunks = select_final_context_chunks(question, ranked_candidates, case_id)
        log_step("ASK_IA", f"Ranking: {len(final_chunks)} chunks selecionados")
        
        # Step 4: Generate response
        answer = generate_response_with_context(question, final_chunks, effective_provider)
        
        # Build response with sources
        sources = [
            {
                "rowid": chunk["rowid"],
                "filename": chunk["filename"],
                "page": chunk["page"],
                "score": round(chunk["rrf_score"], 6)
            }
            for chunk in final_chunks
        ]
        
        response = {
            "answer": answer,
            "hybrid_mode": True,
            "query_expansion": {
                "keywords": expanded["keywords"],
                "semantic_query": expanded["semantic_query"]
            },
            "sources": sources,
            "stats": {
                "fts_results": len(search_results["fts_results"]),
                "vss_results": len(search_results["vss_results"]),
                "final_results": len(final_chunks)
            }
        }
        
        log_step("ASK_IA", f"Resposta gerada: sources={len(response['sources'])}, final_results={response['stats']['final_results']}")
        return response
    
    # Legacy FTS5-only mode
    log_step("ASK_IA", "Modo: LEGACY FTS5-only search")
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    cursor.execute("SELECT content FROM documents_fts WHERE case_id = ? AND documents_fts MATCH ? LIMIT 3", 
                  (case_id, question))
    context = "\n".join([r[0] for r in cursor.fetchall()])
    conn.close()

    prompt = f"Contexto:\n{context}\n\nPergunta: {question}"
    
    if effective_provider == "openrouter":
        return await ask_with_openrouter(prompt)
    elif effective_provider == "gemini":
        url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-1.5-flash:generateContent?key={GEMINI_API_KEY}"
        response = requests.post(url, json={"contents": [{"parts": [{"text": prompt}]}]})
        return response.json()
    
    return {"error": f"Provider {effective_provider} not configured"}


@app.post("/process_pdf")
async def process_pdf(case_id: str = Form(...), file: UploadFile = File(...)):
    """
    Endpoint legacy para processar PDFs (mantido para compatibilidade).
    """
    os.makedirs(UPLOAD_DIR, exist_ok=True)
    file_path = os.path.join(UPLOAD_DIR, file.filename)
    with open(file_path, "wb") as buffer:
        buffer.write(await file.read())
    
    pages = extract_text_from_pdf(file_path)
    store_document(case_id, file.filename, pages)
    
    return {"status": "success", "pages": len(pages), "filename": file.filename}

class DocumentsQuery(BaseModel):
    case_id: str
    top_k: int = 100

class DeleteFileQuery(BaseModel):
    case_id: str
    filename: str

@app.post("/documents")
async def get_documents(query: DocumentsQuery):
    """Get list of documents for a case"""
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    # Get unique filenames with page count for a case
    cursor.execute("""
        SELECT filename, COUNT(*) as page_count, MAX(id) as latest_id
        FROM documents 
        WHERE case_id = ? 
        GROUP BY filename 
        ORDER BY latest_id DESC
        LIMIT ?
    """, (query.case_id, query.top_k))
    results = cursor.fetchall()
    conn.close()
    return [{"filename": r[0], "page_count": r[1]} for r in results]


@app.post("/delete_file")
async def delete_file(request: DeleteFileQuery):
    """Remove a processed document (file + metadata) for a case."""
    filename = os.path.basename(request.filename)
    log_step("DELETE", "Requested document deletion", {"case_id": request.case_id, "filename": filename})

    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    try:
        cursor.execute("SELECT id FROM documents WHERE case_id = ? AND filename = ?", (request.case_id, filename))
        rows = cursor.fetchall()

        if not rows:
            log_step("DELETE", "Document not found for deletion", {"case_id": request.case_id, "filename": filename})
            raise HTTPException(status_code=404, detail="Documento não encontrado")

        doc_ids = [row[0] for row in rows]
        log_step("DELETE", "Deleting database entries", {"ids": doc_ids})

        cursor.executemany("DELETE FROM documents WHERE id = ?", [(doc_id,) for doc_id in doc_ids])
        cursor.executemany("DELETE FROM documents_fts WHERE rowid = ?", [(doc_id,) for doc_id in doc_ids])

        if _vss_available:
            # Verify if vector table exists before attempting delete
            cursor.execute(f"SELECT name FROM sqlite_master WHERE type='table' AND name='{VEC_TABLE}'")
            if cursor.fetchone():
                cursor.executemany(f"DELETE FROM {VEC_TABLE} WHERE rowid = ?", [(doc_id,) for doc_id in doc_ids])

        conn.commit()
        log_step("DELETE", "Database cleanup completed", {"deleted": len(doc_ids)})
    except HTTPException:
        raise
    except Exception as exc:
        log_step("DELETE", "Database error during delete", {"error": str(exc)})
        raise HTTPException(status_code=500, detail=f"Erro ao remover metadados: {exc}")
    finally:
        conn.close()

    file_path = os.path.join(UPLOAD_DIR, filename)
    try:
        if os.path.exists(file_path):
            os.remove(file_path)
            log_step("DELETE", "Arquivo físico removido", {"path": file_path})
        else:
            log_step("DELETE", "Arquivo físico não encontrado (ignorado)", {"path": file_path})
    except Exception as exc:
        log_step("DELETE", "Erro ao remover arquivo físico", {"error": str(exc)})
        raise HTTPException(status_code=500, detail=f"Erro ao remover arquivo: {exc}")

    log_step("DELETE", "Documento excluído com sucesso", {"case_id": request.case_id, "filename": filename})
    return {"status": "success", "message": "Arquivo excluído com sucesso"}

@app.post("/search")
async def search(query: SearchQuery):
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()

    # Build dynamic WHERE clause for metadata filters
    params: list = [query.case_id, query.query]
    extra_conditions = ""

    if query.filename_filter:
        extra_conditions += " AND filename LIKE ?"
        params.append(f"%{query.filename_filter}%")

    if query.file_type_filter:
        ext = query.file_type_filter.lstrip(".")
        extra_conditions += " AND filename LIKE ?"
        params.append(f"%.{ext}")

    params.extend([query.top_k, query.offset])

    cursor.execute(
        f"""SELECT filename, page_number, content,
                   snippet(documents_fts, 0, '<b>', '</b>', '...', 64)
            FROM documents_fts
            WHERE case_id = ? AND documents_fts MATCH ?
            {extra_conditions}
            ORDER BY bm25(documents_fts)
            LIMIT ? OFFSET ?""",
        params,
    )
    results = cursor.fetchall()

    # Total count for pagination metadata (same filters, no LIMIT/OFFSET)
    count_params = [query.case_id, query.query] + params[2:-2]  # strip top_k and offset
    cursor.execute(
        f"""SELECT COUNT(*)
            FROM documents_fts
            WHERE case_id = ? AND documents_fts MATCH ?
            {extra_conditions}""",
        count_params,
    )
    total = cursor.fetchone()[0]
    conn.close()

    return {
        "total": total,
        "offset": query.offset,
        "top_k": query.top_k,
        "results": [{"filename": r[0], "page": r[1], "content": r[2], "snippet": r[3]} for r in results],
    }

    
    return {"error": "Provider not configured"}


async def ask_with_openrouter(prompt: str) -> Dict:
    """
    Faz pergunta à IA usando OpenRouter com fallback automático entre modelos.
    Returns dict with 'answer' field for frontend compatibility.
    """
    for model in OPENROUTER_MODELS:
        try:
            logger.info(f"Trying OpenRouter model: {model}")
            
            response = openrouter_client.chat.completions.create(
                model=model,
                messages=[
                    {"role": "system", "content": "Você é um assistente jurídico helpful. Responda em português brasileiro, de forma clara e profissional."},
                    {"role": "user", "content": prompt}
                ],
                max_tokens=2048,
                temperature=0.7
            )
            
            # Extract answer text from response
            answer_text = response.choices[0].message.content
            logger.info(f"Model {model} succeeded, response length: {len(answer_text)}")
            
            # Return format with 'answer' field for frontend compatibility
            return {
                "answer": answer_text,
                "model": model,
                "provider": "openrouter"
            }
            
        except Exception as e:
            logger.warning(f"Model {model} failed: {str(e)}")
            continue
    
    logger.error("All OpenRouter models failed")
    return {"error": "All models failed", "details": "Unable to get response from any OpenRouter model"}

# ==================== CHUNKED UPLOAD ENDPOINTS ====================

class UploadInitRequest(BaseModel):
    case_id: str
    filename: str
    total_size: int
    total_chunks: int

@app.post("/upload_init")
async def upload_init(request: UploadInitRequest):
    """Initialize a chunked upload session"""
    # Validate file size
    if request.total_size > MAX_FILE_SIZE:
        raise HTTPException(status_code=400, detail=f"File size exceeds maximum limit of {MAX_FILE_SIZE / (1024*1024)}MB")
    
    # Generate unique upload ID
    upload_id = str(uuid.uuid4())
    
    # Create upload directory
    upload_dir = os.path.join(CHUNKS_DIR, upload_id)
    os.makedirs(upload_dir, exist_ok=True)
    
    # Store upload metadata
    metadata = {
        "upload_id": upload_id,
        "case_id": request.case_id,
        "filename": request.filename,
        "total_size": request.total_size,
        "total_chunks": request.total_chunks,
        "uploaded_chunks": [],
        "created_at": time.time()
    }
    
    metadata_file = os.path.join(upload_dir, "metadata.json")
    with open(metadata_file, 'w') as f:
        json.dump(metadata, f)
    
    logger.info(f"Upload initialized: {upload_id} - {request.filename}")
    
    return {
        "status": "success",
        "upload_id": upload_id,
        "chunk_size": CHUNK_SIZE,
        "total_chunks": request.total_chunks
    }

class UploadChunkRequest(BaseModel):
    upload_id: str
    chunk_index: int
    filename: str

@app.post("/upload_chunk")
async def upload_chunk(
    upload_id: str = Form(...),
    chunk_index: int = Form(...),
    filename: str = Form(...),
    case_id: str = Form(...),
    chunk: UploadFile = File(...)
):
    """Receive and store a single chunk"""
    upload_dir = os.path.join(CHUNKS_DIR, upload_id)
    metadata_file = os.path.join(upload_dir, "metadata.json")
    
    # Check if upload session exists
    if not os.path.exists(metadata_file):
        raise HTTPException(status_code=404, detail="Upload session not found or expired")
    
    # Load metadata
    with open(metadata_file, 'r') as f:
        metadata = json.load(f)
    
    # Validate chunk index
    if chunk_index < 0 or chunk_index >= metadata["total_chunks"]:
        raise HTTPException(status_code=400, detail="Invalid chunk index")
    
    # Check if chunk already uploaded
    if chunk_index in metadata["uploaded_chunks"]:
        return {"status": "success", "message": "Chunk already uploaded", "chunk_index": chunk_index}
    
    # Save chunk file
    chunk_filename = f"chunk_{chunk_index:06d}"
    chunk_path = os.path.join(upload_dir, chunk_filename)
    
    with open(chunk_path, "wb") as buffer:
        content = await chunk.read()
        buffer.write(content)
    
    # Update metadata
    metadata["uploaded_chunks"].append(chunk_index)
    metadata["uploaded_chunks"].sort()
    
    with open(metadata_file, 'w') as f:
        json.dump(metadata, f)
    
    progress = (len(metadata["uploaded_chunks"]) / metadata["total_chunks"]) * 100
    
    logger.info(f"Chunk received: {upload_id} - chunk {chunk_index + 1}/{metadata['total_chunks']} ({progress:.1f}%)")
    
    return {
        "status": "success",
        "message": "Chunk uploaded successfully",
        "chunk_index": chunk_index,
        "uploaded_chunks": len(metadata["uploaded_chunks"]),
        "total_chunks": metadata["total_chunks"],
        "progress": progress
    }

@app.get("/upload_status/{upload_id}")
async def get_upload_status(upload_id: str):
    """Get the status of an upload session"""
    upload_dir = os.path.join(CHUNKS_DIR, upload_id)
    metadata_file = os.path.join(upload_dir, "metadata.json")
    
    if not os.path.exists(metadata_file):
        return JSONResponse(status_code=404, content={"status": "error", "message": "Upload session not found or expired"})
    
    with open(metadata_file, 'r') as f:
        metadata = json.load(f)
    
    uploaded_chunks = len(metadata["uploaded_chunks"])
    total_chunks = metadata["total_chunks"]
    
    return {
        "status": "in_progress",
        "upload_id": upload_id,
        "filename": metadata["filename"],
        "uploaded_chunks": metadata["uploaded_chunks"],
        "total_chunks": total_chunks,
        "progress": (uploaded_chunks / total_chunks) * 100 if total_chunks > 0 else 0,
        "missing_chunks": [i for i in range(total_chunks) if i not in metadata["uploaded_chunks"]],
        "created_at": metadata.get("created_at")
    }

class UploadCompleteRequest(BaseModel):
    upload_id: str
    case_id: str

@app.post("/upload_complete")
async def upload_complete(request: UploadCompleteRequest):
    """Finalize a chunked upload by merging all chunks"""
    upload_dir = os.path.join(CHUNKS_DIR, request.upload_id)
    metadata_file = os.path.join(upload_dir, "metadata.json")
    
    if not os.path.exists(metadata_file):
        raise HTTPException(status_code=404, detail="Upload session not found or expired")
    
    with open(metadata_file, 'r') as f:
        metadata = json.load(f)
    
    # Check if all chunks are uploaded
    uploaded_set = set(metadata["uploaded_chunks"])
    expected_set = set(range(metadata["total_chunks"]))
    
    if uploaded_set != expected_set:
        missing = expected_set - uploaded_set
        raise HTTPException(status_code=400, detail=f"Missing chunks: {list(missing)}")
    
    # Merge chunks
    final_path = os.path.join(UPLOAD_DIR, metadata["filename"])
    
    with open(final_path, "wb") as outfile:
        for i in range(metadata["total_chunks"]):
            chunk_path = os.path.join(upload_dir, f"chunk_{i:06d}")
            with open(chunk_path, "rb") as infile:
                outfile.write(infile.read())
    
    # Process the uploaded file
    case_id = request.case_id
    if not case_id:
        raise HTTPException(status_code=400, detail="case_id is required")
    filename = metadata["filename"]
    
    try:
        # Extract text from the document
        pages = extract_text_from_document(final_path)
        
        # Store in database
        store_document(case_id, filename, pages)
        
        logger.info(f"Upload complete: {filename} - {len(pages)} pages processed")
        
    except Exception as e:
        logger.error(f"Error processing uploaded file: {e}")
        # Clean up on error
        if os.path.exists(final_path):
            os.remove(final_path)
        raise HTTPException(status_code=500, detail=f"Error processing file: {str(e)}")
    
    finally:
        # Clean up chunks directory
        shutil.rmtree(upload_dir)
    
    return {
        "status": "success",
        "message": "File uploaded and processed successfully",
        "filename": filename,
        "pages": len(pages)
    }

@app.delete("/upload/{upload_id}")
async def cancel_upload(upload_id: str):
    """Cancel an upload session and clean up chunks"""
    upload_dir = os.path.join(CHUNKS_DIR, upload_id)
    
    if os.path.exists(upload_dir):
        shutil.rmtree(upload_dir)
        logger.info(f"Upload cancelled: {upload_id}")
        return {"status": "success", "message": "Upload cancelled"}
    
    return JSONResponse(status_code=404, content={"status": "error", "message": "Upload session not found"})


# ==================== LAWYER INVITATION ENDPOINTS ====================

# Invitation configuration
INVITATION_EXPIRY_HOURS = 48
MAGIC_LINK_BASE_URL = os.getenv("MAGIC_LINK_BASE_URL", "https://kapjus.kaponline.com.br")

# Password reset token expiry (1 hour)
PASSWORD_RESET_EXPIRY_HOURS = 1


def generate_secure_token(length: int = 64) -> str:
    """Generate a cryptographically secure URL-safe token."""
    return secrets.token_urlsafe(length)


def hash_token(token: str) -> str:
    """Hash a token for secure storage."""
    return hashlib.sha256(token.encode()).hexdigest()


def get_invitation_expiry() -> str:
    """Get expiration datetime for invitation."""
    return (datetime.datetime.utcnow() + datetime.timedelta(hours=INVITATION_EXPIRY_HOURS)).isoformat()


def get_password_reset_expiry() -> str:
    """Get expiration datetime for password reset token."""
    return (datetime.datetime.utcnow() + datetime.timedelta(hours=PASSWORD_RESET_EXPIRY_HOURS)).isoformat()


def init_invitation_tables(conn: sqlite3.Connection):
    """Initialize lawyer invitation tables if they don't exist."""
    cursor = conn.cursor()
    cursor.execute("""
        CREATE TABLE IF NOT EXISTS lawyer_invitations (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            case_id TEXT NOT NULL,
            inviter_email TEXT NOT NULL,
            invitee_email TEXT NOT NULL,
            invitee_name TEXT,
            token TEXT NOT NULL UNIQUE,
            token_hash TEXT NOT NULL,
            role TEXT DEFAULT 'viewer',
            status TEXT DEFAULT 'pending',
            expires_at TEXT NOT NULL,
            created_at TEXT DEFAULT CURRENT_TIMESTAMP,
            accepted_at TEXT,
            revoked_at TEXT,
            access_history TEXT
        )
    """)
    cursor.execute("""
        CREATE TABLE IF NOT EXISTS lawyer_access_logs (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            invitation_id INTEGER NOT NULL,
            lawyer_email TEXT NOT NULL,
            action TEXT NOT NULL,
            ip_address TEXT,
            user_agent TEXT,
            metadata TEXT,
            created_at TEXT DEFAULT CURRENT_TIMESTAMP
        )
    """)
    # Password reset tokens table
    cursor.execute("""
        CREATE TABLE IF NOT EXISTS password_reset_tokens (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            email TEXT NOT NULL,
            token TEXT NOT NULL UNIQUE,
            token_hash TEXT NOT NULL,
            expires_at TEXT NOT NULL,
            used_at TEXT,
            created_at TEXT DEFAULT CURRENT_TIMESTAMP
        )
    """)
    # Create indexes
    cursor.execute("CREATE INDEX IF NOT EXISTS idx_invitations_token ON lawyer_invitations(token)")
    cursor.execute("CREATE INDEX IF NOT EXISTS idx_invitations_email ON lawyer_invitations(invitee_email)")
    cursor.execute("CREATE INDEX IF NOT EXISTS idx_invitations_case ON lawyer_invitations(case_id)")
    cursor.execute("CREATE INDEX IF NOT EXISTS idx_invitations_status ON lawyer_invitations(status)")
    cursor.execute("CREATE INDEX IF NOT EXISTS idx_access_logs_invitation ON lawyer_access_logs(invitation_id)")
    cursor.execute("CREATE INDEX IF NOT EXISTS idx_password_reset_token ON password_reset_tokens(token_hash)")
    cursor.execute("CREATE INDEX IF NOT EXISTS idx_password_reset_email ON password_reset_tokens(email)")
    conn.commit()


class InviteLawyerRequest(BaseModel):
    case_id: str
    inviter_email: str
    inviter_name: Optional[str] = None
    invitee_email: str
    invitee_name: Optional[str] = None
    role: str = "viewer"


class VerifyInvitationRequest(BaseModel):
    token: str
    case_id: str


class RevokeInvitationRequest(BaseModel):
    invitation_id: int
    case_id: str


class ListInvitationsRequest(BaseModel):
    case_id: str


class AccessHistoryRequest(BaseModel):
    case_id: str


class RegisterLawyerPasswordRequest(BaseModel):
    token: str
    case_id: str
    password: str


class ForgotPasswordRequest(BaseModel):
    email: str


class ResetPasswordRequest(BaseModel):
    email: str
    token: str
    password: str


@app.post("/register_lawyer_password")
async def register_lawyer_password(request: RegisterLawyerPasswordRequest):
    """Register password for a lawyer when accepting invitation."""
    token_hash = hash_token(request.token)
    
    conn = sqlite3.connect(DB_PATH)
    init_invitation_tables(conn)
    cursor = conn.cursor()
    
    try:
        # Find the invitation
        cursor.execute("""
            SELECT id, case_id, invitee_email, invitee_name, role, status, expires_at
            FROM lawyer_invitations
            WHERE token_hash = ? AND case_id = ?
        """, (token_hash, request.case_id))
        
        row = cursor.fetchone()
        
        if not row:
            raise HTTPException(status_code=404, detail="Convite não encontrado")
        
        invitation_id, case_id, email, name, role, status, expires_at = row
        
        if status == 'revoked':
            raise HTTPException(status_code=400, detail="Convite foi revogado")
        
        if status == 'accepted':
            raise HTTPException(status_code=400, detail="Convite já foi aceito. Use a página de login.")
        
        if datetime.datetime.fromisoformat(expires_at) < datetime.datetime.utcnow():
            raise HTTPException(status_code=400, detail="Convite expirou")
        
        # Hash the password
        password_hash = hashlib.sha256(request.password.encode()).hexdigest()
        
        # Check if user already exists in users table, update or insert
        cursor.execute("SELECT id FROM users WHERE email = ?", (email,))
        existing_user = cursor.fetchone()
        
        if existing_user:
            # Update existing user password
            cursor.execute("""
                UPDATE users SET password_hash = ? WHERE email = ?
            """, (password_hash, email))
        else:
            # Create new user
            cursor.execute("""
                INSERT INTO users (email, password_hash, name) VALUES (?, ?, ?)
            """, (email, password_hash, name or email.split('@')[0]))
        
        # Update invitation status to accepted
        cursor.execute("""
            UPDATE lawyer_invitations 
            SET status = 'accepted', accepted_at = ?
            WHERE id = ?
        """, (datetime.datetime.utcnow().isoformat(), invitation_id))
        
        conn.commit()
        
        # Send confirmation email with login info
        login_link = f"{MAGIC_LINK_BASE_URL}/login"
        send_password_registered_email(email, name, login_link)
        
        log_step("REGISTER_PWD", f"Password registered for {email}", {"invitation_id": invitation_id})
        
        return {
            "status": "success",
            "message": "Senha cadastrada com sucesso",
            "email": email
        }
        
    except HTTPException:
        raise
    except Exception as e:
        log_step("REGISTER_PWD", f"Error registering password", {"error": str(e)})
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        conn.close()


@app.post("/forgot_password")
async def forgot_password(request: ForgotPasswordRequest):
    """Request password reset - sends email with reset link."""
    email = request.email.lower().strip()
    
    conn = sqlite3.connect(DB_PATH)
    init_invitation_tables(conn)
    cursor = conn.cursor()
    
    try:
        # Check if user exists
        cursor.execute("SELECT id FROM users WHERE email = ?", (email,))
        user = cursor.fetchone()
        
        # Always return success to prevent email enumeration
        # If user exists, send reset email
        if user:
            # Generate reset token
            raw_token = generate_secure_token()
            token_hash = hash_token(raw_token)
            expires_at = get_password_reset_expiry()
            
            # Delete any existing unused tokens for this email
            cursor.execute("""
                DELETE FROM password_reset_tokens 
                WHERE email = ? AND used_at IS NULL
            """, (email,))
            
            # Insert new token
            cursor.execute("""
                INSERT INTO password_reset_tokens (email, token, token_hash, expires_at)
                VALUES (?, ?, ?, ?)
            """, (email, raw_token, token_hash, expires_at))
            
            conn.commit()
            
            # Send reset email
            reset_link = f"{MAGIC_LINK_BASE_URL}/forgot-password?token={raw_token}&email={email}"
            send_password_reset_email(email, reset_link)
            
            log_step("FORGOT_PWD", f"Password reset requested for {email}")
        else:
            log_step("FORGOT_PWD", f"Password reset requested for non-existent email: {email}")
        
        # Always return success
        return {
            "status": "success",
            "message": "Se o e-mail estiver cadastrado, você receberá um link para redefinir sua senha."
        }
        
    except Exception as e:
        log_step("FORGOT_PWD", f"Error requesting password reset", {"error": str(e)})
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        conn.close()


@app.post("/reset_password")
async def reset_password(request: ResetPasswordRequest):
    """Reset password using the reset token."""
    email = request.email.lower().strip()
    token_hash = hash_token(request.token)
    
    conn = sqlite3.connect(DB_PATH)
    init_invitation_tables(conn)
    cursor = conn.cursor()
    
    try:
        # Find valid reset token
        cursor.execute("""
            SELECT id, expires_at, used_at
            FROM password_reset_tokens
            WHERE token_hash = ? AND email = ? AND used_at IS NULL
        """, (token_hash, email))
        
        row = cursor.fetchone()
        
        if not row:
            raise HTTPException(status_code=400, detail="Token de redefinição inválido ou expirado")
        
        token_id, expires_at, used_at = row
        
        if datetime.datetime.fromisoformat(expires_at) < datetime.datetime.utcnow():
            raise HTTPException(status_code=400, detail="Token de redefinição expirou")
        
        # Hash the new password
        password_hash = hashlib.sha256(request.password.encode()).hexdigest()
        
        # Update user password
        cursor.execute("""
            UPDATE users SET password_hash = ? WHERE email = ?
        """, (password_hash, email))
        
        if cursor.rowcount == 0:
            raise HTTPException(status_code=404, detail="Usuário não encontrado")
        
        # Mark token as used
        cursor.execute("""
            UPDATE password_reset_tokens SET used_at = ? WHERE id = ?
        """, (datetime.datetime.utcnow().isoformat(), token_id))
        
        conn.commit()
        
        log_step("RESET_PWD", f"Password reset successful for {email}")
        
        return {
            "status": "success",
            "message": "Senha redefinida com sucesso"
        }
        
    except HTTPException:
        raise
    except Exception as e:
        log_step("RESET_PWD", f"Error resetting password", {"error": str(e)})
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        conn.close()


def send_password_registered_email(to_email: str, name: str = None, login_link: str = "") -> bool:
    """Send confirmation email after password is registered."""
    
    SMTP_HOST = os.getenv("SMTP_HOST", "localhost")
    SMTP_PORT = int(os.getenv("SMTP_PORT", "25"))
    SMTP_FROM = os.getenv("SMTP_FROM", "KapJus <noreply@janeri.com.br>")
    
    display_name = name or to_email.split('@')[0]
    
    html_content = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <style>
        body {{ font-family: 'Segoe UI', Arial, sans-serif; background: #f5f5f5; margin: 0; padding: 20px; }}
        .container {{ max-width: 600px; margin: 0 auto; background: white; border-radius: 12px; overflow: hidden; box-shadow: 0 2px 10px rgba(0,0,0,0.1); }}
        .header {{ background: linear-gradient(135deg, #1a237e 0%, #283593 100%); color: white; padding: 30px; text-align: center; }}
        .header h1 {{ margin: 0; font-size: 28px; }}
        .content {{ padding: 30px; color: #333; line-height: 1.7; }}
        .content h2 {{ color: #1a237e; font-size: 20px; margin-top: 0; }}
        .credentials {{ background: #f8f9fa; border-radius: 8px; padding: 20px; margin: 20px 0; }}
        .credentials p {{ margin: 8px 0; }}
        .credentials strong {{ color: #1a237e; }}
        .button {{ display: inline-block; background: linear-gradient(135deg, #1a237e 0%, #283593 100%); color: white; padding: 16px 32px; text-decoration: none; border-radius: 8px; font-weight: bold; font-size: 16px; margin: 20px 0; }}
        .footer {{ background: #f5f5f5; padding: 20px; text-align: center; font-size: 12px; color: #888; }}
    </style>
</head>
<body>
    <div class="container">
        <div class="header">
            <h1>⚖️ KapJus</h1>
            <p>Sua Inteligência Artificial Jurídica</p>
        </div>
        <div class="content">
            <h2>Olá, {display_name}!</h2>
            <p>Sua conta no <strong>KapJus</strong> foi criada com sucesso!</p>
            
            <div class="credentials">
                <p><strong>🔑 Seus dados de acesso:</strong></p>
                <p><strong>Login (E-mail):</strong> {to_email}</p>
                <p><strong>Senha:</strong> A senha que você acabou de cadastrar</p>
            </div>
            
            <p style="text-align: center;">
                <a href="{login_link}" class="button">Acessar KapJus</a>
            </p>
            
            <p style="font-size: 13px; color: #888;">
                Se você não reconhece esta operação, entre em contato com o administrador.
            </p>
        </div>
        <div class="footer">
            KapJus © 2024 - Inteligência Artificial para o Direito<br>
            Este é um email automático, não responda.
        </div>
    </div>
</body>
</html>"""
    
    text_content = f"""KapJus - Sua Inteligência Jurídica

Olá, {display_name}!

Sua conta no KapJus foi criada com sucesso!

Seus dados de acesso:
- Login (E-mail): {to_email}
- Senha: A senha que você acabou de cadastrar

Acesse aqui: {login_link}

-- 
KapJus © 2024 - Inteligência Artificial para o Direito
"""
    
    try:
        msg = MIMEMultipart('alternative')
        msg['Subject'] = "Sua conta KapJus foi criada com sucesso!"
        msg['From'] = SMTP_FROM
        msg['To'] = to_email
        
        text_part = MIMEText(text_content, 'plain', 'utf-8')
        html_part = MIMEText(html_content, 'html', 'utf-8')
        msg.attach(text_part)
        msg.attach(html_part)
        
        with smtplib.SMTP(SMTP_HOST, SMTP_PORT) as server:
            server.send_message(msg)
        
        log_step("EMAIL", f"Password registered confirmation sent to {to_email}")
        return True
        
    except Exception as e:
        log_step("EMAIL", f"Error sending password registered email to {to_email}: {e}")
        return False


def send_password_reset_email(to_email: str, reset_link: str = "") -> bool:
    """Send password reset email."""
    
    SMTP_HOST = os.getenv("SMTP_HOST", "localhost")
    SMTP_PORT = int(os.getenv("SMTP_PORT", "25"))
    SMTP_FROM = os.getenv("SMTP_FROM", "KapJus <noreply@janeri.com.br>")
    
    html_content = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <style>
        body {{ font-family: 'Segoe UI', Arial, sans-serif; background: #f5f5f5; margin: 0; padding: 20px; }}
        .container {{ max-width: 600px; margin: 0 auto; background: white; border-radius: 12px; overflow: hidden; box-shadow: 0 2px 10px rgba(0,0,0,0.1); }}
        .header {{ background: linear-gradient(135deg, #1a237e 0%, #283593 100%); color: white; padding: 30px; text-align: center; }}
        .header h1 {{ margin: 0; font-size: 28px; }}
        .content {{ padding: 30px; color: #333; line-height: 1.7; }}
        .content h2 {{ color: #1a237e; font-size: 20px; margin-top: 0; }}
        .warning {{ background: #fff3cd; border-left: 4px solid #ffc107; padding: 12px; margin: 20px 0; font-size: 13px; color: #856404; }}
        .button {{ display: inline-block; background: linear-gradient(135deg, #1a237e 0%, #283593 100%); color: white; padding: 16px 32px; text-decoration: none; border-radius: 8px; font-weight: bold; font-size: 16px; margin: 20px 0; }}
        .footer {{ background: #f5f5f5; padding: 20px; text-align: center; font-size: 12px; color: #888; }}
    </style>
</head>
<body>
    <div class="container">
        <div class="header">
            <h1>⚖️ KapJus</h1>
            <p>Sua Inteligência Artificial Jurídica</p>
        </div>
        <div class="content">
            <h2>Recuperação de Senha</h2>
            <p>Você solicitou a recuperação de senha para sua conta no <strong>KapJus</strong>.</p>
            
            <p style="text-align: center;">
                <a href="{reset_link}" class="button">Redefinir Minha Senha</a>
            </p>
            
            <div class="warning">
                ⚠️ Este link expira em 1 hora e é único. Não compartilhe com terceiros.
            </div>
            
            <p style="font-size: 13px; color: #888;">
                Se você não solicitou esta recuperação, ignore este email. Sua senha permanecerá inalterada.
            </p>
        </div>
        <div class="footer">
            KapJus © 2024 - Inteligência Artificial para o Direito<br>
            Este é um email automático, não responda.
        </div>
    </div>
</body>
</html>"""
    
    text_content = f"""KapJus - Recuperação de Senha

Você solicitou a recuperação de senha para sua conta no KapJus.

Acesse aqui: {reset_link}

Este link expira em 1 hora e é único. Não compartilhe com terceiros.

Se você não solicitou esta recuperação, ignore este email.

-- 
KapJus © 2024 - Inteligência Artificial para o Direito
"""
    
    try:
        msg = MIMEMultipart('alternative')
        msg['Subject'] = "Recuperação de Senha - KapJus"
        msg['From'] = SMTP_FROM
        msg['To'] = to_email
        
        text_part = MIMEText(text_content, 'plain', 'utf-8')
        html_part = MIMEText(html_content, 'html', 'utf-8')
        msg.attach(text_part)
        msg.attach(html_part)
        
        with smtplib.SMTP(SMTP_HOST, SMTP_PORT) as server:
            server.send_message(msg)
        
        log_step("EMAIL", f"Password reset email sent to {to_email}")
        return True
        
    except Exception as e:
        log_step("EMAIL", f"Error sending password reset email to {to_email}: {e}")
        return False


@app.post("/invite_lawyer")
async def invite_lawyer(request: InviteLawyerRequest):
    """Send a lawyer invitation with magic link."""
    log_step("INVITE", f"Creating invitation for {request.invitee_email}", {"case_id": request.case_id})
    
    conn = sqlite3.connect(DB_PATH)
    init_invitation_tables(conn)
    cursor = conn.cursor()
    
    try:
        # Check for existing pending invitation
        cursor.execute("""
            SELECT id FROM lawyer_invitations 
            WHERE case_id = ? AND invitee_email = ? AND status = 'pending'
            AND expires_at > datetime('now')
        """, (request.case_id, request.invitee_email))
        existing = cursor.fetchone()
        
        if existing:
            # Return existing invitation
            cursor.execute("""
                SELECT id, token, expires_at FROM lawyer_invitations WHERE id = ?
            """, (existing[0],))
            row = cursor.fetchone()
            conn.close()
            magic_link = f"{MAGIC_LINK_BASE_URL}/magic-login?token={row[1]}&case_id={request.case_id}"
            log_step("INVITE", "Returning existing pending invitation", {"invitation_id": row[0]})
            return {
                "status": "success",
                "message": "Convite pendente já existe",
                "invitation_id": row[0],
                "magic_link": magic_link,
                "expires_at": row[2]
            }
        
        # Generate new token
        raw_token = generate_secure_token()
        token_hash = hash_token(raw_token)
        expires_at = get_invitation_expiry()
        
        # Insert invitation
        cursor.execute("""
            INSERT INTO lawyer_invitations 
            (case_id, inviter_email, invitee_email, invitee_name, token, token_hash, role, status, expires_at)
            VALUES (?, ?, ?, ?, ?, ?, ?, 'pending', ?)
        """, (
            request.case_id, 
            request.inviter_email,
            request.invitee_email,
            request.invitee_name,
            raw_token,
            token_hash,
            request.role,
            expires_at
        ))
        
        invitation_id = cursor.lastrowid
        conn.commit()
        
        # Fetch case info for the email
        case_name = ""
        case_description = ""
        try:
            cursor.execute("SELECT name, description FROM cases WHERE id = ?", (request.case_id,))
            case_row = cursor.fetchone()
            if case_row:
                case_name = case_row[0] or ""
                case_description = case_row[1] or ""
        except Exception as e:
            log_step("INVITE", f"Could not fetch case info: {e}")
        
        magic_link = f"{MAGIC_LINK_BASE_URL}/magic-login?token={raw_token}&case_id={request.case_id}"
        
        # Send invitation email
        inviter_display_name = request.inviter_name if request.inviter_name else request.inviter_email
        email_sent = send_invitation_email(request.invitee_email, request.case_id, magic_link, inviter_display_name, case_name, case_description)
        
        log_step("INVITE", f"Invitation created", {"invitation_id": invitation_id, "expires_at": expires_at})
        
        return {
            "status": "success",
            "message": "Convite enviado com sucesso",
            "invitation_id": invitation_id,
            "magic_link": magic_link,
            "expires_at": expires_at
        }
        
    except Exception as e:
        log_step("INVITE", f"Error creating invitation", {"error": str(e)})
        raise HTTPException(status_code=500, detail=f"Erro ao criar convite: {str(e)}")
    finally:
        conn.close()


@app.post("/verify_invitation")
async def verify_invitation(request: VerifyInvitationRequest):
    """Verify a magic link token and return invitation status."""
    token_hash = hash_token(request.token)
    
    conn = sqlite3.connect(DB_PATH)
    init_invitation_tables(conn)
    cursor = conn.cursor()
    
    try:
        cursor.execute("""
            SELECT id, case_id, invitee_email, invitee_name, role, status, expires_at
            FROM lawyer_invitations
            WHERE token_hash = ? AND case_id = ?
        """, (token_hash, request.case_id))
        
        row = cursor.fetchone()
        
        if not row:
            raise HTTPException(status_code=404, detail="Convite não encontrado")
        
        invitation_id, case_id, email, name, role, status, expires_at = row
        
        if status == 'revoked':
            raise HTTPException(status_code=400, detail="Convite foi revogado")
        
        if status == 'accepted':
            raise HTTPException(status_code=400, detail="Convite já foi aceito")
        
        if datetime.datetime.fromisoformat(expires_at) < datetime.datetime.utcnow():
            raise HTTPException(status_code=400, detail="Convite expirou")
        
        # Generate session token for the lawyer
        session_token = generate_secure_token(32)
        session_hash = hash_token(session_token)
        
        # Update invitation with accepted status
        cursor.execute("""
            UPDATE lawyer_invitations 
            SET status = 'accepted', accepted_at = ?
            WHERE id = ?
        """, (datetime.datetime.utcnow().isoformat(), invitation_id))
        
        conn.commit()
        
        log_step("VERIFY", f"Invitation verified and accepted", {"invitation_id": invitation_id})
        
        return {
            "status": "valid",
            "invitation_id": invitation_id,
            "case_id": case_id,
            "lawyer_email": email,
            "lawyer_name": name,
            "role": role,
            "session_token": session_token,
            "expires_at": expires_at
        }
        
    except HTTPException:
        raise
    except Exception as e:
        log_step("VERIFY", f"Error verifying invitation", {"error": str(e)})
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        conn.close()


@app.post("/revoke_invitation")
async def revoke_invitation(request: RevokeInvitationRequest):
    """Revoke a pending invitation."""
    conn = sqlite3.connect(DB_PATH)
    init_invitation_tables(conn)
    cursor = conn.cursor()
    
    try:
        cursor.execute("""
            UPDATE lawyer_invitations 
            SET status = 'revoked', revoked_at = ?
            WHERE id = ? AND case_id = ? AND status = 'pending'
        """, (datetime.datetime.utcnow().isoformat(), request.invitation_id, request.case_id))
        
        if cursor.rowcount == 0:
            raise HTTPException(status_code=404, detail="Convite não encontrado ou já processado")
        
        conn.commit()
        
        log_step("REVOKE", f"Invitation revoked", {"invitation_id": request.invitation_id})
        
        return {"status": "success", "message": "Convite revogado com sucesso"}
        
    except HTTPException:
        raise
    except Exception as e:
        log_step("REVOKE", f"Error revoking invitation", {"error": str(e)})
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        conn.close()


@app.post("/invitations")
async def list_invitations(request: ListInvitationsRequest):
    """List all invitations for a case."""
    conn = sqlite3.connect(DB_PATH)
    init_invitation_tables(conn)
    cursor = conn.cursor()
    
    try:
        cursor.execute("""
            SELECT id, invitee_email, invitee_name, role, status, 
                   created_at, expires_at, accepted_at
            FROM lawyer_invitations
            WHERE case_id = ?
            ORDER BY created_at DESC
        """, (request.case_id,))
        
        rows = cursor.fetchall()
        
        invitations = []
        for row in rows:
            invitations.append({
                "id": row[0],
                "invitee_email": row[1],
                "invitee_name": row[2],
                "role": row[3],
                "status": row[4],
                "created_at": row[5],
                "expires_at": row[6],
                "accepted_at": row[7]
            })
        
        return {"status": "success", "invitations": invitations}
        
    except Exception as e:
        log_step("INVITATIONS", f"Error listing invitations", {"error": str(e)})
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        conn.close()


@app.post("/access_history")
async def access_history(request: AccessHistoryRequest):
    """Get access history for a case."""
    conn = sqlite3.connect(DB_PATH)
    init_invitation_tables(conn)
    cursor = conn.cursor()
    
    try:
        cursor.execute("""
            SELECT al.id, al.lawyer_email, al.action, al.ip_address,
                   al.user_agent, al.created_at, li.invitee_name
            FROM lawyer_access_logs al
            LEFT JOIN lawyer_invitations li ON al.invitation_id = li.id
            WHERE li.case_id = ?
            ORDER BY al.created_at DESC
            LIMIT 100
        """, (request.case_id,))
        
        rows = cursor.fetchall()
        
        logs = []
        for row in rows:
            logs.append({
                "id": row[0],
                "lawyer_email": row[1],
                "lawyer_name": row[6] if row[6] else row[1],
                "action": row[2],
                "ip_address": row[3],
                "user_agent": row[4],
                "accessed_at": row[5]
            })
        
        return {"status": "success", "logs": logs}
        
    except Exception as e:
        log_step("ACCESS_HISTORY", f"Error getting access history", {"error": str(e)})
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        conn.close()


@app.get("/health")
async def health_check():
    """Health check endpoint for service monitoring"""
    return {"status": "healthy", "service": "kapjus-rag"}

if __name__ == "__main__":
    socket_dir = os.path.dirname(SOCKET_PATH)
    os.makedirs(socket_dir, exist_ok=True)
    if os.path.exists(SOCKET_PATH):
        os.remove(SOCKET_PATH)
    
    def run_uvicorn_tcp():
        """Run uvicorn on TCP port 8000"""
        uvicorn.run(app, host="127.0.0.1", port=8000, log_level="info")
    
    # Start TCP server in background thread
    tcp_thread = threading.Thread(target=run_uvicorn_tcp, daemon=True)
    tcp_thread.start()
    
    # Run main server on Unix socket
    uvicorn.run(app, uds=SOCKET_PATH, log_level="info")
